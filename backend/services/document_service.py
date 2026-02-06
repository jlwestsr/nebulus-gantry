"""
Document Service for managing the Knowledge Vault.

Provides:
- Collection CRUD operations
- Document upload, processing, and indexing
- Text extraction from PDF, DOCX, TXT, CSV
- Chunking and embedding into ChromaDB
- Semantic search across documents
"""
import io
import logging
from typing import Optional

import chromadb
from sqlalchemy.orm import Session as DBSession

from backend.config import Settings
from backend.models.collection import Collection
from backend.models.document import Document

logger = logging.getLogger(__name__)

# Chunk settings: ~500 tokens ≈ 2000 chars, with 100 char overlap
CHUNK_SIZE = 2000
CHUNK_OVERLAP = 100


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise ValueError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content."""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(content))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_txt(content: bytes) -> str:
    """Extract text from TXT/CSV content."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("latin-1")
        except Exception as e:
            logger.error(f"Failed to decode text file: {e}")
            raise ValueError(f"Failed to decode text file: {e}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks."""
    if not text:
        return []

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # Try to break at a sentence or paragraph boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind("\n\n", start, end)
            if para_break > start + chunk_size // 2:
                end = para_break + 2
            else:
                # Look for sentence break
                sentence_break = text.rfind(". ", start, end)
                if sentence_break > start + chunk_size // 2:
                    end = sentence_break + 2

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Move start with overlap
        start = end - overlap if end < len(text) else len(text)

    return chunks


class DocumentService:
    """Service for managing documents and collections in the Knowledge Vault."""

    def __init__(self, db: DBSession):
        self.db = db
        self.settings = Settings()
        self._chroma_client: Optional[chromadb.HttpClient] = None
        self._chroma_available = False

        try:
            host_url = self.settings.chroma_host
            if host_url.startswith("http://"):
                host_url = host_url[7:]
            elif host_url.startswith("https://"):
                host_url = host_url[8:]

            if ":" in host_url:
                host, port_str = host_url.split(":", 1)
                port = int(port_str)
            else:
                host = host_url
                port = 8000

            self._chroma_client = chromadb.HttpClient(host=host, port=port)
            self._chroma_available = True
            logger.info("DocumentService: ChromaDB connected")
        except Exception as e:
            logger.warning(f"ChromaDB unavailable for documents: {e}")

    def _get_collection_name(self, user_id: int) -> str:
        """Get the ChromaDB collection name for a user's documents."""
        return f"user_{user_id}_documents"

    def _get_chroma_collection(self, user_id: int):
        """Get or create the ChromaDB collection for a user."""
        if not self._chroma_available or not self._chroma_client:
            return None
        try:
            return self._chroma_client.get_or_create_collection(
                name=self._get_collection_name(user_id)
            )
        except Exception as e:
            logger.error(f"Failed to get ChromaDB collection: {e}")
            return None

    # ========== Collection CRUD ==========

    def create_collection(
        self, user_id: int, name: str, description: str | None = None
    ) -> Collection:
        """Create a new collection."""
        collection = Collection(
            user_id=user_id,
            name=name,
            description=description,
        )
        self.db.add(collection)
        self.db.commit()
        self.db.refresh(collection)
        return collection

    def list_collections(self, user_id: int) -> list[Collection]:
        """List all collections for a user."""
        return (
            self.db.query(Collection)
            .filter(Collection.user_id == user_id)
            .order_by(Collection.created_at.desc())
            .all()
        )

    def get_collection(self, collection_id: int, user_id: int) -> Collection | None:
        """Get a single collection by ID."""
        return (
            self.db.query(Collection)
            .filter(Collection.id == collection_id, Collection.user_id == user_id)
            .first()
        )

    def update_collection(
        self,
        collection_id: int,
        user_id: int,
        name: str | None = None,
        description: str | None = None,
    ) -> Collection | None:
        """Update a collection."""
        collection = self.get_collection(collection_id, user_id)
        if not collection:
            return None

        if name is not None:
            collection.name = name
        if description is not None:
            collection.description = description

        self.db.commit()
        self.db.refresh(collection)
        return collection

    def delete_collection(self, collection_id: int, user_id: int) -> bool:
        """Delete a collection and all its documents."""
        collection = self.get_collection(collection_id, user_id)
        if not collection:
            return False

        # Delete document chunks from ChromaDB
        chroma_collection = self._get_chroma_collection(user_id)
        if chroma_collection:
            for doc in collection.documents:
                try:
                    # Delete all chunks for this document
                    ids_to_delete = [
                        f"doc_{doc.id}_chunk_{i}" for i in range(doc.chunk_count)
                    ]
                    if ids_to_delete:
                        chroma_collection.delete(ids=ids_to_delete)
                except Exception as e:
                    logger.warning(f"Failed to delete document chunks from ChromaDB: {e}")

        self.db.delete(collection)
        self.db.commit()
        return True

    # ========== Document CRUD ==========

    def upload_document(
        self,
        user_id: int,
        filename: str,
        content: bytes,
        content_type: str,
        collection_id: int | None = None,
    ) -> Document:
        """Upload and process a document."""
        # Validate collection if provided
        if collection_id is not None:
            collection = self.get_collection(collection_id, user_id)
            if not collection:
                raise ValueError("Collection not found")

        # Create document record
        document = Document(
            user_id=user_id,
            collection_id=collection_id,
            filename=filename,
            content_type=content_type,
            file_size=len(content),
            status="processing",
        )
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)

        try:
            # Extract text based on content type
            if content_type == "pdf":
                text = extract_text_from_pdf(content)
            elif content_type == "docx":
                text = extract_text_from_docx(content)
            elif content_type in ("txt", "csv"):
                text = extract_text_from_txt(content)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

            if not text.strip():
                raise ValueError("No text could be extracted from the document")

            # Chunk the text
            chunks = chunk_text(text)
            document.chunk_count = len(chunks)

            # Index chunks in ChromaDB
            chroma_collection = self._get_chroma_collection(user_id)
            if chroma_collection and chunks:
                ids = [f"doc_{document.id}_chunk_{i}" for i in range(len(chunks))]
                metadatas = [
                    {
                        "document_id": document.id,
                        "chunk_index": i,
                        "filename": filename,
                        "collection_id": collection_id or -1,
                    }
                    for i in range(len(chunks))
                ]
                chroma_collection.add(ids=ids, documents=chunks, metadatas=metadatas)

            document.status = "ready"
            self.db.commit()
            self.db.refresh(document)

        except Exception as e:
            logger.error(f"Failed to process document: {e}")
            document.status = "failed"
            document.error_message = str(e)
            self.db.commit()
            self.db.refresh(document)

        return document

    def list_documents(
        self, user_id: int, collection_id: int | None = None
    ) -> list[Document]:
        """List documents for a user, optionally filtered by collection."""
        query = self.db.query(Document).filter(Document.user_id == user_id)
        if collection_id is not None:
            query = query.filter(Document.collection_id == collection_id)
        return query.order_by(Document.created_at.desc()).all()

    def get_document(self, document_id: int, user_id: int) -> Document | None:
        """Get a single document by ID."""
        return (
            self.db.query(Document)
            .filter(Document.id == document_id, Document.user_id == user_id)
            .first()
        )

    def delete_document(self, document_id: int, user_id: int) -> bool:
        """Delete a document and its chunks from ChromaDB."""
        document = self.get_document(document_id, user_id)
        if not document:
            return False

        # Delete chunks from ChromaDB
        chroma_collection = self._get_chroma_collection(user_id)
        if chroma_collection and document.chunk_count > 0:
            try:
                ids_to_delete = [
                    f"doc_{document_id}_chunk_{i}"
                    for i in range(document.chunk_count)
                ]
                chroma_collection.delete(ids=ids_to_delete)
            except Exception as e:
                logger.warning(f"Failed to delete chunks from ChromaDB: {e}")

        self.db.delete(document)
        self.db.commit()
        return True

    # ========== Search ==========

    def search_documents(
        self,
        user_id: int,
        query: str,
        collection_ids: list[int] | None = None,
        top_k: int = 5,
    ) -> list[dict]:
        """Search documents using semantic search."""
        chroma_collection = self._get_chroma_collection(user_id)
        if not chroma_collection:
            return []

        try:
            # Build filter for collection_ids if provided
            where_filter = None
            if collection_ids:
                # ChromaDB uses $in for list matching
                where_filter = {"collection_id": {"$in": collection_ids}}

            results = chroma_collection.query(
                query_texts=[query],
                n_results=top_k,
                where=where_filter,
            )

            formatted_results = []
            if results["documents"] and results["documents"][0]:
                documents = results["documents"][0]
                distances = (
                    results["distances"][0]
                    if results.get("distances")
                    else [0] * len(documents)
                )
                metadatas = (
                    results["metadatas"][0]
                    if results.get("metadatas")
                    else [{}] * len(documents)
                )

                for i, doc in enumerate(documents):
                    meta = metadatas[i] if i < len(metadatas) else {}
                    # Convert distance to similarity (lower distance = higher similarity)
                    similarity = 1.0 / (1.0 + (distances[i] if i < len(distances) else 0))
                    formatted_results.append({
                        "document_id": meta.get("document_id", 0),
                        "filename": meta.get("filename", "Unknown"),
                        "chunk_text": doc,
                        "similarity": similarity,
                        "chunk_index": meta.get("chunk_index", 0),
                    })

            return formatted_results

        except Exception as e:
            logger.error(f"Failed to search documents: {e}")
            return []

    def build_rag_context(
        self,
        user_id: int,
        query: str,
        document_scope: list[dict] | None = None,
        top_k: int = 3,
    ) -> str:
        """Build RAG context from document search results.

        Args:
            user_id: The user ID
            query: The search query
            document_scope: Optional list of {"type": "document"|"collection", "id": int}
            top_k: Number of results to return

        Returns:
            Formatted context string with source citations
        """
        # Parse document scope into collection_ids and document_ids
        collection_ids = []
        document_ids = []

        if document_scope:
            for scope in document_scope:
                if scope.get("type") == "collection":
                    collection_ids.append(scope["id"])
                elif scope.get("type") == "document":
                    document_ids.append(scope["id"])

        # Search with collection filter if any
        results = self.search_documents(
            user_id=user_id,
            query=query,
            collection_ids=collection_ids if collection_ids else None,
            top_k=top_k * 2,  # Get more to filter by document_ids
        )

        # Filter by document_ids if specified
        if document_ids:
            results = [r for r in results if r["document_id"] in document_ids]

        # Limit to top_k
        results = results[:top_k]

        if not results:
            return ""

        # Format as context with citations
        parts = ["Relevant document context:"]
        for r in results:
            citation = f"[Source: {r['filename']}, chunk {r['chunk_index'] + 1}]"
            parts.append(f"- {r['chunk_text'][:500]}... {citation}")

        return "\n".join(parts)
